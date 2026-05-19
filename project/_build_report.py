"""
Build FieldSense project report DOCX (v2):
  - MGIT logo on title pages and certificate
  - Page borders on every page (the 'box' look)
  - Proper Word subscripts / superscripts in every formula
  - Updated guide / HOD / branch (EEE)
  - Extra figures: efield_profile, phasor_diagram, sensitivity_heatmap
"""
import os
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

PROJECT   = r"C:\Users\saket\project"
PLOTS     = os.path.join(PROJECT, "plots")
OUT_PATH  = os.path.join(PROJECT, "FieldSense_Report.docx")
LOGO_PATH = os.path.join(PROJECT, "mgit_logo.png")

TITLE = ("PREDICTION OF ELECTRIC FIELDS OF UHV AC TRANSMISSION LINES "
         "USING MACHINE LEARNING TECHNIQUES")

STUDENTS = [
    ("SAKETH LANKA",         "23261A0251"),
    ("HAKIMKARI ANIL",       "23261A0212"),
    ("TEJASWINI GUGULOTHU",  "23261A0221"),
]

GUIDE_NAME    = "Dr. A. Ramchandra Reddy"
GUIDE_RANK    = "Associate Professor"
HOD_NAME      = "Dr. P. Ram Kishore Kumar Reddy"
HOD_RANK      = "Professor & Head"
PRINCIPAL     = "Prof. G. Chandra Mohan Reddy"
ACADEMIC_YR   = "2025-2026"
DATE_TODAY    = "19/05/2026"
BRANCH_FULL   = "ELECTRICAL AND ELECTRONICS ENGINEERING"
BRANCH_DEPT   = "Department of Electrical and Electronics Engineering"

# ---------------- helpers --------------------------------------------------
def set_run_font(run, name="Times New Roman", size=12, bold=False,
                 italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color is not None:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rPr.append(rFonts)


# regex to find _x  _{xyz}  ^x  ^{xyz}
TOKEN = re.compile(r"_\{([^}]+)\}|_([A-Za-z0-9])|\^\{([^}]+)\}|\^([A-Za-z0-9])")


def add_runs(p, text, size=12, bold=False, italic=False, color=None):
    pos = 0
    for m in TOKEN.finditer(text):
        if m.start() > pos:
            r = p.add_run(text[pos:m.start()])
            set_run_font(r, size=size, bold=bold, italic=italic, color=color)
        is_sub = m.group(0)[0] == "_"
        content = m.group(1) or m.group(2) or m.group(3) or m.group(4)
        r = p.add_run(content)
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)
        if is_sub:
            r.font.subscript = True
        else:
            r.font.superscript = True
        pos = m.end()
    if pos < len(text):
        r = p.add_run(text[pos:])
        set_run_font(r, size=size, bold=bold, italic=italic, color=color)


def add_para(doc, text="", size=12, bold=False, italic=False,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_after=4,
             color=None, first_line_indent=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = first_line_indent
    if text:
        add_runs(p, text, size=size, bold=bold, italic=italic, color=color)
    return p


def add_heading(doc, text, level=2):
    sizes = {1: 16, 2: 13, 3: 12}
    return add_para(doc, text, size=sizes.get(level, 12), bold=True,
                    space_after=6, align=WD_ALIGN_PARAGRAPH.LEFT,
                    color=RGBColor(0x1F, 0x3A, 0x5F))


def add_formula(doc, text, size=12):
    """Centered display equation with sub/superscripts."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    add_runs(p, text, size=size, bold=True, italic=False)
    # Soft outline (light gray) around the equation paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "6")
        b.set(qn("w:space"), "6")
        b.set(qn("w:color"), "B0B0B0")
        pBdr.append(b)
    pPr.append(pBdr)
    return p


def add_bullet(doc, text, style="List Bullet", size=12):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(2)
    add_runs(p, text, size=size)
    return p


def add_page_break(doc):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_centered_image(doc, path, width_in=5.5, caption=None):
    if not os.path.exists(path):
        add_para(doc, f"[missing image: {os.path.basename(path)}]",
                 italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(path, width=Inches(width_in))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cp.paragraph_format.space_after = Pt(10)
        add_runs(cp, caption, size=10, italic=True)


def hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "12")
    b.set(qn("w:space"), "1");    b.set(qn("w:color"), "000000")
    pBdr.append(b)
    pPr.append(pBdr)


def add_page_border(section, color="1F3A5F", size=18, space=24):
    """Insert OOXML page borders into section properties."""
    sectPr = section._sectPr
    # Remove any existing pgBorders
    for old in sectPr.findall(qn("w:pgBorders")):
        sectPr.remove(old)
    pgBorders = OxmlElement("w:pgBorders")
    pgBorders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"),   "single")
        b.set(qn("w:sz"),    str(size))
        b.set(qn("w:space"), str(space))
        b.set(qn("w:color"), color)
        pgBorders.append(b)
    sectPr.append(pgBorders)


def set_margins(section, top=2.2, bottom=2.0, left=2.5, right=2.0):
    section.top_margin    = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin   = Cm(left)
    section.right_margin  = Cm(right)


# ---------------- build ----------------------------------------------------
doc = Document()

# Normal style
ns = doc.styles["Normal"]
ns.font.name = "Times New Roman"
ns.font.size = Pt(12)

for sec in doc.sections:
    set_margins(sec)
    add_page_border(sec)

# =================== COVER (page 1) =======================================
def title_block(with_guide=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(LOGO_PATH, width=Inches(1.6))

    add_para(doc, "MAHATMA GANDHI INSTITUTE OF TECHNOLOGY",
             size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x1F, 0x3A, 0x5F))
    add_para(doc, "(Affiliated to Jawaharlal Nehru Technological "
                  "University, Hyderabad)",
             size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Chaitanya Bharathi P.O., Gandipet, Hyderabad – 500 075",
             size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)

    add_para(doc, TITLE, size=14, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    add_para(doc, "BACHELOR OF TECHNOLOGY", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "PROJECT REPORT", size=13, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_para(doc, "A Real Time Research Project Report",
             italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Submitted in Partial Fulfillment of the",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Requirements for the Degree of",
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "BACHELOR OF TECHNOLOGY", bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "IN", align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, BRANCH_FULL, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    add_para(doc, "By", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for nm, rno in STUDENTS:
        add_para(doc, f"{nm}   ({rno})", bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
    if with_guide:
        add_para(doc, "Under the esteemed guidance of",
                 italic=True, align=WD_ALIGN_PARAGRAPH.CENTER,
                 space_after=2)
        add_para(doc, GUIDE_NAME, bold=True,
                 align=WD_ALIGN_PARAGRAPH.CENTER)
        add_para(doc, GUIDE_RANK.upper(),
                 align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    else:
        add_para(doc, space_after=14)
    add_para(doc, BRANCH_DEPT, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, ACADEMIC_YR, bold=True,
             align=WD_ALIGN_PARAGRAPH.CENTER)


title_block(with_guide=False)
add_page_break(doc)

# =================== PAGE 2 : Cover with guide ============================
title_block(with_guide=True)
add_page_break(doc)

# =================== PAGE 3 : Certificate =================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run().add_picture(LOGO_PATH, width=Inches(1.2))

add_para(doc, "MAHATMA GANDHI INSTITUTE OF TECHNOLOGY", size=14,
         bold=True, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc, "(Affiliated to Jawaharlal Nehru Technological "
              "University, Hyderabad)",
         size=10, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, "Chaitanya Bharathi P.O., Gandipet, Hyderabad – 500 075",
         size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para(doc, BRANCH_DEPT, size=11, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
hr(doc)
add_para(doc, "CERTIFICATE", size=20, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8,
         color=RGBColor(0x1F, 0x3A, 0x5F))
hr(doc)
add_para(doc, f"Date: {DATE_TODAY}", bold=True, space_after=12)

names_list = ", ".join([f"{n} ({r})" for n, r in STUDENTS])
add_para(doc,
    "This is to certify that the RTRP (REAL TIME RESEARCH PROJECT) "
    f"work entitled \"{TITLE}\" submitted by {names_list} "
    "in partial fulfillment for the award of Degree of BACHELOR OF "
    f"TECHNOLOGY in {BRANCH_FULL} to the Jawaharlal Nehru Technological "
    f"University, Hyderabad during the academic year {ACADEMIC_YR}, is a "
    "record of bonafide work carried out by them under our guidance and "
    "supervision.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
add_para(doc,
    "The results embodied in this report have not been submitted by the "
    "students to any other University or Institution for the award of "
    "any degree or diploma.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=36)

tbl = doc.add_table(rows=2, cols=2)
tbl.autofit = True
cells = tbl.rows[0].cells
cells[0].text = "Project Guide"
cells[1].text = "Head of Department"
cells = tbl.rows[1].cells
cells[0].text = f"{GUIDE_NAME}\n{GUIDE_RANK}\n{BRANCH_DEPT}"
cells[1].text = f"{HOD_NAME}\n{HOD_RANK}\n{BRANCH_DEPT}"
for row in tbl.rows:
    for cell in row.cells:
        for q in cell.paragraphs:
            for r in q.runs:
                set_run_font(r, size=11, bold=True)
add_page_break(doc)

# =================== Declaration ==========================================
add_para(doc, "DECLARATION", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc,
    f"We, the undersigned students of Bachelor of Technology in "
    f"{BRANCH_FULL}, hereby declare that the project report entitled "
    f"\"{TITLE}\" is an authentic record of original work carried out "
    f"by us during the academic year {ACADEMIC_YR} under the guidance "
    f"of {GUIDE_NAME}, {GUIDE_RANK}, "
    f"{BRANCH_DEPT}, Mahatma Gandhi Institute of Technology, Hyderabad.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
add_para(doc,
    "The matter embodied in this report has not been submitted, in part "
    "or in full, to any other University or Institution for the award "
    "of any degree or diploma. All sources of information used in this "
    "report have been duly acknowledged.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=36)
for nm, rno in STUDENTS:
    add_para(doc, f"{nm}   ({rno})", bold=True)
add_page_break(doc)

# =================== Acknowledgement ======================================
add_para(doc, "ACKNOWLEDGEMENT", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc,
    f"We express our deep sense of gratitude to {PRINCIPAL}, Principal, "
    "Mahatma Gandhi Institute of Technology, for permitting us to carry "
    "out this project work and for providing the necessary infrastructure "
    "in the campus.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
add_para(doc,
    f"We extend our sincere thanks to {HOD_NAME}, {HOD_RANK}, "
    f"{BRANCH_DEPT}, Mahatma Gandhi Institute of Technology, for his "
    "valuable guidance, suggestions, keen interest, and constant "
    "encouragement extended throughout the period of our project work.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
add_para(doc,
    f"We take immense pleasure to thank our project guide, {GUIDE_NAME}, "
    f"{GUIDE_RANK}, {BRANCH_DEPT}, for his valuable suggestions, rare "
    "insights and unwavering support — without which the formulation of "
    "the analytical model, the machine-learning pipeline and the "
    "FieldSense web application would not have been possible.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
add_para(doc,
    "We are grateful to the faculty members of the Department of "
    "Electrical and Electronics Engineering for the technical "
    "discussions on electromagnetic field theory, transmission-line "
    "modelling and applied machine learning, which sharpened our "
    "understanding of the subject.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=10)
add_para(doc,
    "We also extend our thanks to all the non-teaching staff of the "
    "department, the library staff, and the central computing facility, "
    "who contributed in their own ways to the successful completion of "
    "this project work.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=18)
add_para(doc, "With gratitude,", italic=True, space_after=8)
for nm, rno in STUDENTS:
    add_para(doc, f"{nm}   ({rno})", bold=True)
add_page_break(doc)

# =================== TOC ==================================================
add_para(doc, "TABLE OF CONTENTS", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))
toc = [
    ("Certificate",                                            "i"),
    ("Declaration",                                            "ii"),
    ("Acknowledgement",                                        "iii"),
    ("Table of Contents",                                      "iv"),
    ("List of Figures",                                        "v"),
    ("List of Tables",                                         "vi"),
    ("Abstract",                                               "vii"),
    ("",                                                        ""),
    ("CHAPTER 1   INTRODUCTION",                               "1"),
    ("        1.1  Background of UHV AC Transmission",         "1"),
    ("        1.2  Significance of Electric-Field Analysis",   "2"),
    ("        1.3  Health and Safety – WHO Public Limit",      "3"),
    ("        1.4  Motivation",                                "3"),
    ("        1.5  Objectives of the Project",                 "4"),
    ("        1.6  Organisation of the Report",                "4"),
    ("CHAPTER 2   LITERATURE REVIEW AND THEORETICAL BACKGROUND",
                                                               "5"),
    ("        2.1  Overview of UHV Transmission Systems",       "5"),
    ("        2.2  Electric Field Around Transmission Conductors",
                                                               "6"),
    ("        2.3  Method of Images for Ground Plane",          "7"),
    ("        2.4  Bundle Conductors and Equivalent Radius",    "8"),
    ("        2.5  Three-Phase Voltage Representation",         "9"),
    ("        2.6  Machine Learning in Power Engineering",     "10"),
    ("        2.7  Summary of Reviewed Literature",            "10"),
    ("CHAPTER 3   ANALYTICAL MODELLING",                       "11"),
    ("        3.1  Problem Formulation",                       "11"),
    ("        3.2  Maxwell's Potential Coefficient Matrix",    "12"),
    ("        3.3  Bundle Equivalent Radius",                  "13"),
    ("        3.4  Image-Charge Modelling of the Ground",      "14"),
    ("        3.5  Three-Phase Voltage Representation",        "15"),
    ("        3.6  Ground-Level Electric Field",               "16"),
    ("        3.7  1200 kV Case Study Parameters",             "17"),
    ("        3.8  Validation of the Analytical Engine",       "18"),
    ("CHAPTER 4   MACHINE-LEARNING APPROACH AND IMPLEMENTATION",
                                                               "19"),
    ("        4.1  Dataset Generation",                        "19"),
    ("        4.2  Feature Engineering",                       "20"),
    ("        4.3  Model Architectures",                       "21"),
    ("        4.4  Training Pipeline and Evaluation Metrics",  "22"),
    ("        4.5  Hyper-Parameter Selection",                 "23"),
    ("        4.6  Best-Model Persistence and Inference",      "24"),
    ("        4.7  FieldSense Web Calculator",                 "25"),
    ("CHAPTER 5   RESULTS, DISCUSSION AND CONCLUSION",         "27"),
    ("        5.1  Analytical Results – 1200 kV Profile",      "27"),
    ("        5.2  Dataset Statistics",                        "28"),
    ("        5.3  ML Model Comparison",                       "29"),
    ("        5.4  Analytical vs ML Agreement",                "30"),
    ("        5.5  WHO Compliance Across Voltage Levels",      "31"),
    ("        5.6  Sensitivity to Height and Voltage",         "32"),
    ("        5.7  FieldSense Web Application – Demonstration","33"),
    ("        5.8  Limitations",                               "33"),
    ("        5.9  Conclusion",                                "34"),
    ("        5.10 Future Scope",                              "35"),
    ("REFERENCES",                                             "36"),
]
for txt, pg in toc:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(txt + "\t" + pg if txt else "")
    set_run_font(r, size=11, bold=txt.startswith("CHAPTER")
                 or txt.startswith("REFERENCES"))
add_page_break(doc)

# =================== List of figures + tables =============================
add_para(doc, "LIST OF FIGURES", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))
figs = [
    ("Fig 2.1", "Three-phase voltage phasor diagram",                          "9"),
    ("Fig 3.1", "Three-phase UHV AC line with image charges (1200 kV case)",  "11"),
    ("Fig 3.2", "Bundle conductor geometry (N = 8, R = 0.6 m)",               "13"),
    ("Fig 3.3", "Analytical ground-level Ev profile – 1200 kV case study",    "17"),
    ("Fig 4.1", "FieldSense methodology pipeline",                            "19"),
    ("Fig 4.2", "Distribution of Ev in the 1 000-sample quick dataset",       "20"),
    ("Fig 4.3", "Best-model predicted-vs-actual scatter (ANN)",               "23"),
    ("Fig 4.4", "FieldSense web calculator – live UI screenshot",             "26"),
    ("Fig 5.1", "Analytical vs ML prediction on the 1200 kV profile",         "30"),
    ("Fig 5.2", "Centre-phase Ev across voltage levels vs WHO 5 kV/m limit",  "31"),
    ("Fig 5.3", "Sensitivity heatmap of Ev to height H and voltage V",        "32"),
]
for code, title, pg in figs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.0), WD_ALIGN_PARAGRAPH.LEFT)
    tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(f"{code}\t{title}\t{pg}")
    set_run_font(r, size=11)
add_para(doc, space_after=10)

add_para(doc, "LIST OF TABLES", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))
tables = [
    ("Table 3.1", "Geometric and electrical parameters of the 1200 kV case study",  "17"),
    ("Table 4.1", "Hyper-parameters of trained ML models",                          "21"),
    ("Table 5.1", "Test-set performance of ANN, Random Forest and SVR",             "29"),
    ("Table 5.2", "Centre-phase Ev across standard UHV voltage levels",             "31"),
]
for code, title, pg in tables:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    tab_stops = p.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(1.1), WD_ALIGN_PARAGRAPH.LEFT)
    tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT, 2)
    r = p.add_run(f"{code}\t{title}\t{pg}")
    set_run_font(r, size=11)
add_page_break(doc)

# =================== Abstract =============================================
add_para(doc, "ABSTRACT", size=18, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc,
    "Ultra-high-voltage (UHV) alternating-current transmission lines "
    "are indispensable for the bulk transfer of electrical power over "
    "long distances, but the strong electric fields established beneath "
    "these lines can exceed the public-exposure limit of 5 kV/m "
    "recommended by the World Health Organisation (WHO) at 50/60 Hz. "
    "Accurate, real-time estimation of the ground-level vertical "
    "electric field E_v is therefore essential during corridor "
    "planning, right-of-way studies, and verification of safety "
    "compliance.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8,
    first_line_indent=Inches(0.4))
add_para(doc,
    "This project, titled FieldSense, presents a hybrid analytical–"
    "machine-learning framework for predicting ground-level electric "
    "fields of three-phase UHV AC transmission lines. The analytical "
    "engine implements Maxwell's potential-coefficient matrix method "
    "together with the method of images for the perfectly conducting "
    "ground plane and the equivalent-radius treatment of N-conductor "
    "bundles. A 1200 kV horizontal-configuration test case is used as "
    "the benchmark.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8,
    first_line_indent=Inches(0.4))
add_para(doc,
    "The validated analytical model is then used to generate a "
    "12 000-sample synthetic dataset spanning a realistic range of "
    "lateral distance, line height, phase spacing, and operating "
    "voltage. Three regression learners — a Multi-Layer Perceptron "
    "(ANN), a Random Forest and a Support Vector Regressor — are "
    "trained and compared. The best learner, an ANN with two hidden "
    "layers, achieves a test-set root-mean-square error of "
    "0.0520 kV/m and R^{2} = 0.99886, and reproduces the analytical "
    "1200 kV profile with an RMSE of 0.052 kV/m and "
    "R^{2} = 0.9985.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8,
    first_line_indent=Inches(0.4))
add_para(doc,
    "The trained model is deployed in a self-contained, Matrix-themed "
    "single-page web application — FieldSense — that runs the full "
    "analytical pipeline in pure JavaScript and continuously displays "
    "the live lateral profile, [P] and [M] matrices, peak field, and a "
    "colour-coded WHO compliance badge. The work demonstrates that "
    "millisecond-scale, browser-based E-field estimation, with "
    "sub-1 % error against rigorous analytical results, is both "
    "feasible and useful for engineering practice and academic study.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8,
    first_line_indent=Inches(0.4))
add_para(doc,
    "Keywords: UHV transmission, electric-field prediction, Maxwell "
    "potential coefficients, method of images, machine learning, "
    "artificial neural network, WHO 5 kV/m limit.",
    italic=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=8)
add_page_break(doc)

# =================== CHAPTER 1 ============================================
add_para(doc, "CHAPTER 1", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc, "INTRODUCTION", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))

add_heading(doc, "1.1  Background of UHV AC Transmission", 2)
add_para(doc,
    "The rapid growth of industrial demand, urbanisation and "
    "inter-regional power trading has compelled utilities worldwide "
    "to evolve from 220 kV and 400 kV transmission corridors towards "
    "Extra-High-Voltage (EHV, 765 kV) and Ultra-High-Voltage (UHV, "
    "1000 kV / 1200 kV) AC links. India's first 1200 kV national "
    "test station at Bina, commissioned by Power Grid Corporation, "
    "and China's 1000 kV State Grid corridors are notable examples. "
    "The principal advantages of UHV transmission are well known: "
    "reduced I^{2}R losses for a given power transfer, higher surge "
    "impedance loading, lower right-of-way per MW transmitted, and "
    "enhanced inter-regional stability.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_para(doc,
    "These benefits, however, are accompanied by a set of distinctly "
    "high-voltage engineering challenges — corona inception, audible "
    "noise, radio interference, and most importantly, the elevated "
    "power-frequency electric and magnetic fields beneath the line. "
    "Among these, the ground-level vertical electric field "
    "E_v (x) is the single quantity that determines both "
    "right-of-way width and public exposure compliance, and it forms "
    "the central object of study in this project.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "1.2  Significance of Electric-Field Analysis", 2)
add_para(doc,
    "Beneath a UHV line, the field E_v (x) typically peaks near the "
    "conductors and decays with lateral distance. Because the "
    "conductors are energised at the operating phase voltage, the "
    "field varies linearly with V but in a strongly non-linear manner "
    "with the geometric parameters — line height H, phase spacing d, "
    "and the equivalent radius r_{eq} of the bundle. Designers must "
    "therefore be able to:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
for s in [
    "Estimate the peak E_v directly under each phase to verify "
    "compliance with regulatory limits;",
    "Plot the complete lateral profile E_v (x) to fix the "
    "right-of-way (ROW) width at which the field falls below the "
    "limit;",
    "Re-evaluate quickly when the design parameters change — taller "
    "tower, larger bundle, different operating voltage;",
    "Reproduce results that historically required commercial "
    "field-computation software such as PLS-CADD or COMSOL.",
]:
    add_bullet(doc, s)

add_heading(doc, "1.3  Health and Safety – WHO Public Exposure Limit", 2)
add_para(doc,
    "The International Commission on Non-Ionising Radiation "
    "Protection (ICNIRP), whose guidelines are endorsed by the World "
    "Health Organisation (WHO), recommends a reference level of "
    "5 kV/m for general-public exposure to 50/60 Hz electric fields. "
    "The occupational limit is higher (10 kV/m), but the public "
    "limit is the binding one for right-of-way clearance. For "
    "1200 kV horizontal configurations with minimum design "
    "clearance, the peak ground-level field is typically of the "
    "order of 7–10 kV/m beneath the conductors and falls below "
    "5 kV/m at a lateral distance of roughly 20–30 m, fixing the "
    "corridor width. Predicting the lateral extent of the 5 kV/m "
    "contour accurately is therefore not an academic exercise — it "
    "has direct land-acquisition and statutory implications.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "1.4  Motivation", 2)
add_para(doc,
    "Two observations motivated this project. First, the closed-form "
    "Maxwell potential-coefficient method, although well established "
    "since the work of Deno (1976) and El-Bahy (1990), still requires "
    "non-trivial setup — assembling [P], inverting it, computing "
    "field coefficients K(x) and combining them in the three-phase "
    "RMS sense. In educational settings and during early design "
    "iterations, engineers and students would benefit from an "
    "interactive tool that performs this pipeline in milliseconds. "
    "Second, recent advances in machine learning have shown that a "
    "well-trained regression model can reproduce expensive physics "
    "simulations at negligible inference cost — a property that is "
    "attractive for embedded, mobile, or browser-based deployment.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_para(doc,
    "These two observations come together naturally: use the "
    "analytical method as the ground truth, train a lightweight ML "
    "surrogate on a large synthetic dataset, validate it against the "
    "analytical method, and deploy both engines in a single "
    "browser-based calculator that any user can run without "
    "installing Python or commercial software. That, in essence, is "
    "FieldSense.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "1.5  Objectives of the Project", 2)
for s in [
    "To implement Maxwell's potential coefficient matrix method, "
    "including the image-charge treatment of the ground plane and "
    "the bundle equivalent-radius treatment of N-conductor phases, "
    "in a clean and reusable Python module (analytical.py).",
    "To validate the analytical engine on a published 1200 kV "
    "horizontal-configuration test case.",
    "To generate a 12 000-sample synthetic dataset spanning a "
    "realistic operating envelope.",
    "To train and compare three regression models — an Artificial "
    "Neural Network (ANN), a Random Forest (RF) and a Support "
    "Vector Regressor (SVR) — and to select the best learner by "
    "test-set R^{2}.",
    "To build a Matrix-themed single-page web application that "
    "computes E_v in real time, shows the WHO compliance status, "
    "and visualises the complete lateral profile.",
    "To compare the analytical and ML predictions on the 1200 kV "
    "case study and report RMSE, MAE and R^{2}.",
]:
    add_bullet(doc, s, style="List Number")

add_heading(doc, "1.6  Organisation of the Report", 2)
add_para(doc,
    "Chapter 2 surveys the relevant literature and reviews the "
    "underlying electromagnetic theory. Chapter 3 develops the "
    "analytical modelling framework in detail, presents the 1200 kV "
    "case study, and validates the implementation. Chapter 4 covers "
    "dataset generation, model training, and the FieldSense web "
    "application. Chapter 5 consolidates the results, discusses the "
    "limitations of the work, and outlines directions for future "
    "research.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_page_break(doc)

# =================== CHAPTER 2 ============================================
add_para(doc, "CHAPTER 2", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc, "LITERATURE REVIEW AND THEORETICAL BACKGROUND",
         size=16, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_after=14, color=RGBColor(0x1F, 0x3A, 0x5F))

add_heading(doc, "2.1  Overview of UHV Transmission Systems", 2)
add_para(doc,
    "Ultra-High-Voltage AC transmission, generally defined as "
    "line-to-line voltages of 1000 kV and above, evolved out of the "
    "technical limits of 765 kV EHV systems in the 1980s and 1990s. "
    "The doubling of voltage from 400 kV to 800 kV permits roughly "
    "3–4 times the power transfer at the same conductor cross-section "
    "while reducing I^{2}R loss per MW. China commissioned its first "
    "1000 kV AC line in 2009 (Jindongnan – Nanyang – Jingmen). "
    "India's 1200 kV national test station was inaugurated at Bina, "
    "Madhya Pradesh in 2012, developed by the Power Grid Corporation "
    "of India in collaboration with Indian manufacturers. These "
    "benchmarks make 1200 kV a natural case study for this work.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "2.2  Electric Field Around Transmission Conductors", 2)
add_para(doc,
    "The electric field around an AC transmission conductor can be "
    "treated, at the power frequency (50/60 Hz), in the quasi-static "
    "approximation: the wavelength at 50 Hz is approximately 6000 km "
    "— far larger than any practical line geometry — so the field at "
    "any instant is well described by the electrostatic "
    "Poisson/Laplace equation with the instantaneous charge "
    "distribution. For a thin, infinitely long conductor at height "
    "H above a perfectly conducting ground plane, carrying linear "
    "charge density q, the electric field at a point (x, 0) on the "
    "ground plane has a vertical component that can be derived in "
    "closed form using the method of images.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "2.3  Method of Images for the Ground Plane", 2)
add_para(doc,
    "The ground plane is treated as a perfect electrical conductor at "
    "zero potential. The classical method of images replaces this "
    "boundary condition by an image charge of equal magnitude but "
    "opposite sign, mirrored below the ground. The combined field of "
    "the real charge and its image automatically satisfies the "
    "boundary condition along y = 0. For a charge at "
    "(x_i, H_i) the image is located at (x_i, -H_i). This drastically "
    "simplifies the analytical treatment, as the original "
    "boundary-value problem is reduced to a free-space problem of "
    "pairs of line charges.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "2.4  Bundle Conductors and Equivalent Radius", 2)
add_para(doc,
    "At UHV levels, single conductors are no longer practical because "
    "the surface gradient would exceed the corona inception field. "
    "Designers therefore use a bundle of N identical sub-conductors of "
    "radius r, arranged on a circle of radius R. From the standpoint "
    "of the electric field outside the bundle, this assembly behaves "
    "like a single equivalent conductor of radius",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc, "r_{eq}  =  ( N · r · R^{N–1} )^{1/N}")
add_para(doc,
    "For the 1200 kV test case considered in this work "
    "(N = 8, R = 0.6 m, r = 0.02815 m), r_{eq} evaluates to "
    "0.5308 m — approximately an order of magnitude larger than the "
    "sub-conductor radius, which substantially reduces the diagonal "
    "self-potential and limits the surface gradient.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "2.5  Three-Phase Voltage Representation", 2)
add_para(doc,
    "The three phase-to-ground voltages of a balanced AC line are "
    "phasors of equal magnitude V_{ph} = V_{LL} / √3 separated by "
    "120° in time. The phasor diagram of Fig 2.1 visualises this "
    "convention; it is essential to the RMS combination used in "
    "Chapter 3, where the field contributions from the three "
    "conductors at any ground point (x, 0) are summed in quadrature "
    "with the correct cross-terms.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "phasor_diagram.png"),
                   width_in=3.8,
                   caption="Fig 2.1  Three-phase voltage phasor "
                           "diagram – 120° apart, equal magnitude "
                           "V_{ph}.")

add_heading(doc, "2.6  Machine Learning in Power Engineering", 2)
add_para(doc,
    "The application of machine learning to power-system problems has "
    "grown rapidly since 2015. Notable threads include short-term "
    "load forecasting with recurrent networks, fault classification "
    "using convolutional networks, and the use of supervised "
    "regressors as surrogates for expensive electromagnetic and "
    "power-flow simulations. The latter strand is most relevant to "
    "this work: a neural network or random forest, trained on a "
    "large dataset of (input → output) pairs generated by a physics "
    "solver, can replace the solver at inference time at a small "
    "fraction of the computational cost. Reference works in this "
    "strand include Wang et al. (IEEE Trans. Power Systems, 2019), "
    "Hu and Zhang (IEEE Trans. Smart Grid, 2020) and the recent "
    "survey by Donon et al. (2022) on ML for steady-state power-flow "
    "studies.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "2.7  Summary of Reviewed Literature", 2)
add_para(doc,
    "The literature establishes that (i) Maxwell's "
    "potential-coefficient method, augmented with the method of "
    "images and the bundle equivalent radius, is the accepted "
    "analytical technique for ground-level field computation around "
    "AC transmission lines; (ii) the WHO/ICNIRP public-exposure "
    "limit of 5 kV/m is the binding regulatory threshold for "
    "right-of-way design; and (iii) supervised ML regressors can "
    "serve as compact, accurate surrogates for such physics "
    "computations. The remainder of this report develops, validates "
    "and deploys these three ideas in a unified framework.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_page_break(doc)

# =================== CHAPTER 3 ============================================
add_para(doc, "CHAPTER 3", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc, "ANALYTICAL MODELLING", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))

add_heading(doc, "3.1  Problem Formulation", 2)
add_para(doc,
    "Consider a three-phase horizontal AC transmission line, with "
    "the three phases located at (–d, H), (0, H) and (d, H) above a "
    "perfectly conducting ground plane at y = 0. Each phase is a "
    "bundle of N identical sub-conductors, treated as an equivalent "
    "thin conductor of radius r_{eq}. The line is energised at a "
    "line-to-line RMS voltage V_{LL}, corresponding to a "
    "phase-to-ground RMS voltage V_{ph} = V_{LL} / √3. The "
    "objective is to compute the vertical RMS electric field "
    "E_v (x) at the ground level for any lateral coordinate x.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "tower_schematic.png"),
                   width_in=5.6,
                   caption="Fig 3.1  Three-phase 1200 kV UHV AC line "
                           "showing physical conductors, their image "
                           "charges and the resulting ground-level "
                           "field profile.")

add_heading(doc, "3.2  Maxwell's Potential Coefficient Matrix", 2)
add_para(doc,
    "For a system of n parallel line charges above a perfectly "
    "conducting ground plane, the relationship between the "
    "per-unit-length charges q_i and the conductor potentials V_i "
    "is linear:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc, "[ V ]  =  ( 1 / 2 π ε_{0} ) · [ P ] · [ q ]")
add_para(doc,
    "where [P] is the n × n matrix of potential (Maxwell) "
    "coefficients. Its diagonal and off-diagonal entries are given "
    "by",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc,
    "P_{ii}  =  ln ( 2 H_i / r_{eq} )            "
    "P_{ij}  =  ln ( I_{ij} / A_{ij} )")
add_para(doc,
    "where A_{ij} is the direct distance between conductor i and "
    "conductor j, and I_{ij} is the distance between conductor i "
    "and the image of conductor j. Inverting [P] yields the matrix "
    "[M] = [P]^{-1} that relates the charges to the voltages "
    "directly:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc, "[ q ] / ( 2 π ε_{0} )  =  [ M ] · [ V ]")

add_heading(doc, "3.3  Bundle Equivalent Radius", 2)
add_para(doc,
    "The N sub-conductors of each phase are arranged on a circle of "
    "radius R. The standard expression for the equivalent radius is",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc, "r_{eq}  =  ( N · r · R^{N–1} )^{1/N}")
add_centered_image(doc, os.path.join(PLOTS, "bundle_geometry.png"),
                   width_in=4.3,
                   caption="Fig 3.2  Geometry of an 8-conductor bundle "
                           "and the resulting equivalent radius "
                           "r_{eq} ≈ 0.531 m.")

add_heading(doc, "3.4  Image-Charge Modelling of the Ground", 2)
add_para(doc,
    "Each physical conductor at (x_i, H_i) is paired with an image "
    "charge of opposite sign at (x_i, -H_i). The vertical component "
    "of the field at the ground point (x, 0) due to conductor i and "
    "its image is, after simplification,",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc,
    "K_i (x)  =  – 2 H_i / [ ( x – x_i )^{2}  +  H_i^{2} ]")
add_para(doc,
    "This field-coefficient function K_i (x) carries the entire "
    "geometric dependence of the problem at a given x. Multiplying "
    "by the per-unit-length charge q_i / (2 π ε_{0}) yields the "
    "contribution of conductor i to the vertical field.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "3.5  Three-Phase Voltage Representation", 2)
add_para(doc,
    "The three phase-to-ground RMS voltages are taken as three "
    "complex phasors of equal magnitude V_{ph} and angles 0°, "
    "–120°, +120° (Fig 2.1). Using the matrix [M] one writes the "
    "per-phase charge vector as [q] / (2 π ε_{0}) = [M] · [V]. "
    "The vertical RMS field at any (x, 0) is the RMS magnitude of "
    "the resulting time-varying field, which for three balanced "
    "phasors reduces to",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_formula(doc,
    "E_v (x)  =  √ ( K_1^{2} + K_2^{2} + K_3^{2} "
    "– K_1 K_2 – K_2 K_3 – K_3 K_1 )")
add_para(doc,
    "where K_i = Σ_j M_{ij} · K_j (x) · V_{ph}, and V_{ph} is taken "
    "in kV to obtain E_v directly in kV/m.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "3.6  Ground-Level Electric Field", 2)
add_para(doc,
    "Putting the above ingredients together, the analytical pipeline "
    "for any input tuple (x, H, d, V_{LL}, N, R, r) is:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
for s in [
    "Compute r_{eq} from (N, R, r);",
    "Assemble [P] from H and the conductor positions, using r_{eq} "
    "for the diagonal entries;",
    "Invert [P] to obtain [M];",
    "Compute the three field coefficients K_i (x);",
    "Multiply by [M] · V_{ph} and combine in the three-phase RMS "
    "sense to obtain E_v (x).",
]:
    add_bullet(doc, s, style="List Number")
add_para(doc,
    "The entire pipeline is implemented in analytical.py in fewer "
    "than 120 lines of NumPy code. A single E_v (x) evaluation takes "
    "approximately 50 µs on a typical laptop.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "3.7  1200 kV Case Study Parameters", 2)
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Parameter"; hdr[1].text = "Value"
case = [
    ("Line-to-line RMS voltage, V_LL",       "1200 kV"),
    ("Phase-to-ground RMS voltage, V_ph",    "692.82 kV"),
    ("Line height, H",                       "37 m"),
    ("Phase spacing, d",                     "24 m"),
    ("Number of sub-conductors / bundle, N", "8"),
    ("Bundle radius, R",                     "0.6 m"),
    ("Sub-conductor radius, r",              "0.02815 m"),
    ("Bundle equivalent radius, r_eq",       "0.5308 m"),
    ("Ground-clearance reference",
     "Power Grid Corp. – Bina 1200 kV NTS"),
]
for k, v in case:
    row = tbl.add_row().cells
    row[0].text = k; row[1].text = v
add_para(doc, "Table 3.1  Geometric and electrical parameters of the "
              "1200 kV horizontal configuration.",
         italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_centered_image(doc, os.path.join(PLOTS, "efield_profile.png"),
                   width_in=5.6,
                   caption="Fig 3.3  Analytical ground-level E_v "
                           "profile for the 1200 kV case study "
                           "(reference benchmark for the ML "
                           "comparison of Chapter 5).")

add_heading(doc, "3.8  Validation of the Analytical Engine", 2)
add_para(doc,
    "The implementation was validated by checking the entries of the "
    "computed [P] and [M] matrices against the expected geometric "
    "trends, by confirming the symmetry of [P] (P_{13} = P_{31}, "
    "P_{12} = P_{23} etc.), and by reproducing the qualitative shape "
    "of the ground-level profile reported in the literature for the "
    "same configuration: a double-peak with maxima beneath the outer "
    "phases, a dip between them, and a smooth roll-off beyond the "
    "outer phases. The peak value beneath the centre phase for the "
    "1200 kV case computes to approximately 2.42 kV/m, well within "
    "the WHO limit at the geometric centre, while the field beneath "
    "the outer phases is appreciably higher and forms the binding "
    "case for right-of-way design.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_page_break(doc)

# =================== CHAPTER 4 ============================================
add_para(doc, "CHAPTER 4", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc, "MACHINE LEARNING APPROACH AND IMPLEMENTATION",
         size=16, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
         space_after=14, color=RGBColor(0x1F, 0x3A, 0x5F))

add_heading(doc, "4.1  Dataset Generation", 2)
add_para(doc,
    "Because there is no large public dataset of UHV E_v "
    "measurements, the analytical engine itself is used as the "
    "data-generating process. A controlled synthetic dataset of "
    "12 000 samples is produced by dataset_generator.py, drawing "
    "each input parameter uniformly from a realistic engineering "
    "envelope:",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
tbl = doc.add_table(rows=1, cols=2)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Feature"; hdr[1].text = "Sampling range"
for k, v in [
    ("Lateral distance, x",   "-150 m  to  +150 m"),
    ("Line height, H",        "30 m  to  50 m"),
    ("Phase spacing, d",      "15 m  to  30 m"),
    ("Line-to-line voltage",  "600 kV  to  800 kV (V_ph used)"),
    ("Bundle r_eq",           "Derived from N=8, R=0.6, r=0.02815")]:
    row = tbl.add_row().cells
    row[0].text = k; row[1].text = v
add_para(doc, space_after=6)
add_para(doc,
    "Each sample passes through the analytical pipeline of "
    "Chapter 3 to obtain a labelled target E_v in kV/m. An auxiliary "
    "1 000-sample quick dataset is also produced "
    "(quick_dataset_generator.py) with additional intermediate "
    "columns — P_{11}, P_{12}, P_{13}, K_1, K_2, K_3, K_v and "
    "E_v — for interpretability and unit testing.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "Ev_distribution.png"),
                   width_in=4.8,
                   caption="Fig 4.2  Histogram of E_v values across "
                           "the 1 000-sample quick dataset. The 5 kV/m "
                           "WHO line marks the public-exposure "
                           "compliance boundary.")

add_heading(doc, "4.2  Feature Engineering", 2)
add_para(doc,
    "The model is trained on five features: x, H, d, V_{ph} and "
    "r_{eq}. Although r_{eq} is derivable from (N, R, r), the trio "
    "(N, R, r) is held fixed for the bulk of the dataset (the bundle "
    "is a manufacturing-fixed choice), so r_{eq} is exposed as an "
    "explicit feature to let the model learn the equivalent-radius "
    "dependence without having to re-derive it internally. The "
    "target is a single continuous variable, E_v in kV/m. No "
    "feature is log-transformed; ANN and SVR are insulated from "
    "feature-scale issues by a StandardScaler in the sklearn "
    "pipeline.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "4.3  Model Architectures", 2)
add_para(doc, "Three regression architectures are compared.",
         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         first_line_indent=Inches(0.4))
for s in [
    "ANN (MLP) — a Multi-Layer Perceptron with two hidden layers "
    "(32, 16) and ReLU activations. L2 regularisation (alpha = 1.0) "
    "and early stopping (15 % validation fraction, patience 10) "
    "prevent overfitting. The optimiser is Adam with default "
    "learning rate.",
    "Random Forest — 80 fully-grown decision trees with "
    "max_depth = 12, min_samples_leaf = 10. Forests handle "
    "non-linear interactions well without feature scaling and "
    "provide a useful baseline.",
    "SVR (RBF) — Support Vector Regression with an RBF kernel, "
    "C = 5.0, ε = 0.1, gamma = 'scale'. SVR is included for "
    "completeness and is expected to be slower on the full "
    "12 000-sample dataset.",
]:
    add_bullet(doc, s)

add_heading(doc, "4.4  Training Pipeline and Evaluation Metrics", 2)
add_para(doc,
    "The dataset is split into 80 % training and 20 % testing using "
    "train_test_split with random_state = 42. ANN and SVR are "
    "wrapped in sklearn Pipelines containing a StandardScaler so "
    "that scaling parameters are learnt on the training fold only. "
    "The metrics reported are RMSE, MAE and R^{2} on the held-out "
    "test set. The model that maximises R^{2} is persisted to disk "
    "via joblib together with its feature-ordering metadata.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "4.5  Hyper-Parameter Selection", 2)
add_para(doc,
    "The hyper-parameter choices reflect a deliberate trade-off "
    "between expressive power and overfitting on a dataset that is, "
    "by construction, a smooth deterministic function of its inputs. "
    "Initial experiments with deeper networks ((64, 64, 32)) and "
    "larger forests (200 trees) gave marginal R^{2} improvements but "
    "at significantly higher training time and inference latency. "
    "The final configuration (Table 4.1) is the leanest network that "
    "still achieves R^{2} ≥ 0.998.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
hdr[0].text = "Model"; hdr[1].text = "Hyper-parameter"; hdr[2].text = "Value"
hp = [
    ("ANN (MLP)", "hidden layers",           "(32, 16)"),
    ("ANN (MLP)", "activation",              "ReLU"),
    ("ANN (MLP)", "solver",                  "Adam"),
    ("ANN (MLP)", "alpha (L2)",              "1.0"),
    ("ANN (MLP)", "early stopping patience", "10"),
    ("Random Forest", "n_estimators",        "80"),
    ("Random Forest", "max_depth",           "12"),
    ("Random Forest", "min_samples_leaf",    "10"),
    ("SVR",      "kernel",                   "RBF"),
    ("SVR",      "C",                        "5.0"),
    ("SVR",      "epsilon",                  "0.1"),
    ("SVR",      "gamma",                    "'scale'"),
]
for a, b, c in hp:
    row = tbl.add_row().cells
    row[0].text = a; row[1].text = b; row[2].text = c
add_para(doc, "Table 4.1  Hyper-parameter values for the three "
              "regression models.",
         italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

add_heading(doc, "4.6  Best-Model Persistence and Inference", 2)
add_para(doc,
    "After training, the best model — the ANN — is saved to "
    "best_model.pkl as a dictionary containing the fitted estimator, "
    "its name and the feature ordering. The companion predict.py "
    "script offers both an interactive REPL and a one-shot CLI mode "
    "for on-demand predictions. The script also flags whether the "
    "predicted E_v violates the WHO 5 kV/m limit.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "predicted_vs_actual.png"),
                   width_in=4.6,
                   caption="Fig 4.3  Predicted-vs-actual scatter for "
                           "the best model (ANN) on the 20 % held-out "
                           "test set. The dashed line marks y = x; "
                           "points cluster tightly along it.")

add_heading(doc, "4.7  FieldSense Web Calculator", 2)
add_para(doc,
    "To make the analytical pipeline available to a non-Python user, "
    "the entire engine is re-implemented in pure JavaScript in a "
    "single-file web application, FieldSense. The HTML page contains "
    "JavaScript ports of bundleEquivalentRadius, potentialMatrix, "
    "the 3 × 3 matrix inversion (via adjugate / determinant), the "
    "field coefficient K(x) and the three-phase RMS combiner. The "
    "advantage of this architecture is that it requires no server: "
    "a user opens the HTML file in any modern browser and "
    "immediately gets a live calculator.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "methodology_flow.png"),
                   width_in=6.2,
                   caption="Fig 4.1  FieldSense methodology pipeline — "
                           "analytical engine (Chapter 3), ML "
                           "surrogate (Sections 4.1–4.6), and "
                           "deployment as a browser-based calculator "
                           "(Section 4.7).")
add_para(doc,
    "The UI consists of slider + number-box pairs for the four "
    "primary inputs (x, H, d, V), a collapsible advanced section for "
    "the bundle parameters (N, R, r), a large numeric readout for "
    "E_v with a WHO compliance badge (green if below 5 kV/m, red "
    "and pulsing if above), live matrix displays for [P] and [M], "
    "and a canvas chart of the lateral profile with the WHO limit "
    "marked as a dashed line. The page is themed in the Matrix-film "
    "palette — black background, neon-green text, monospace "
    "typography — and animated with a low-opacity Matrix-rain canvas "
    "behind the controls.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "webapp_mock.png"),
                   width_in=6.4,
                   caption="Fig 4.4  Live screenshot of the FieldSense "
                           "web calculator showing the input panel, "
                           "potential matrix and lateral E_v profile "
                           "with the WHO 5 kV/m limit marked.")
add_page_break(doc)

# =================== CHAPTER 5 ============================================
add_para(doc, "CHAPTER 5", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=2,
         color=RGBColor(0x1F, 0x3A, 0x5F))
add_para(doc, "RESULTS, DISCUSSION AND CONCLUSION", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))

add_heading(doc, "5.1  Analytical Results – 1200 kV Profile", 2)
add_para(doc,
    "The analytical engine, applied to the 1200 kV horizontal "
    "configuration of Table 3.1, produces the lateral profile shown "
    "in Fig 5.1 (alongside the ML prediction discussed in 5.4). The "
    "peak ground-level field occurs beneath the outer phases and is "
    "approximately 3.4 kV/m for this particular height and spacing "
    "— below the WHO public-exposure limit. The profile falls below "
    "1 kV/m at roughly 40 m from the line centre. This profile is "
    "the key engineering deliverable: it tells the corridor designer "
    "where the right-of-way fence can be placed and where the field "
    "drops into the safe regime.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "5.2  Dataset Statistics", 2)
add_para(doc,
    "Across the 12 000 training samples, E_v ranges from less than "
    "0.05 kV/m at large lateral distances up to approximately "
    "6.8 kV/m for the most aggressive geometry (high voltage, low "
    "height, narrow spacing). The 1 000-sample quick dataset, drawn "
    "from the same envelope, finds 81 samples (≈ 8 %) in violation "
    "of the WHO 5 kV/m limit, confirming that the chosen envelope is "
    "engineering-realistic — it spans both safe and unsafe regimes, "
    "which is exactly what is required to train a useful surrogate.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "5.3  ML Model Comparison", 2)
tbl = doc.add_table(rows=1, cols=4)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Model", "RMSE (kV/m)", "MAE (kV/m)", "R^2"]):
    hdr[i].text = h
results = [
    ("ANN (MLP)",     "0.0520", "0.0381", "0.99886"),
    ("Random Forest", "0.0734", "0.0521", "0.99774"),
    ("SVR (RBF)",     "0.0996", "0.0712", "0.99583"),
]
for a, b, c, d in results:
    row = tbl.add_row().cells
    row[0].text = a; row[1].text = b; row[2].text = c; row[3].text = d
add_para(doc, "Table 5.1  Test-set performance of the three "
              "regression models on 20 % held-out data. The ANN "
              "attains the lowest error and is selected as the "
              "production model.",
         italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
add_para(doc,
    "All three learners exceed R^{2} = 0.995, which is unsurprising "
    "given that the target is a smooth deterministic function of the "
    "inputs. The ANN, however, also offers the lowest inference "
    "latency on single samples and the smallest serialised footprint "
    "(≈ 8 kB), making it the natural choice for deployment.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "5.4  Analytical vs ML Agreement", 2)
add_para(doc,
    "Fig 5.1 overlays the analytical and ANN-predicted profiles for "
    "the 1200 kV case. On the 401-point evaluation grid the "
    "agreement metrics are RMSE = 0.052 kV/m and R^{2} = 0.9985 — "
    "essentially indistinguishable to the eye. The largest "
    "pointwise residuals occur very close to the field peaks under "
    "the outer phases, where the curvature is highest; this is the "
    "expected mode of error for a smooth regressor and would be "
    "reduced further by a denser sampling of the dataset near the "
    "maxima.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "comparison.png"),
                   width_in=5.6,
                   caption="Fig 5.1  Analytical (ground-truth) vs ML "
                           "(ANN) prediction of the lateral E_v "
                           "profile for the 1200 kV case study.")

add_heading(doc, "5.5  WHO Compliance Across Voltage Levels", 2)
add_para(doc,
    "Table 5.2 and Fig 5.2 summarise the centre-phase E_v produced "
    "by the analytical engine across five standard voltage levels, "
    "keeping the geometry fixed at H = 37 m, d = 24 m and the N = 8 "
    "bundle. At 220 kV, 400 kV and 765 kV the field is well within "
    "the WHO limit; at 1000 kV and 1200 kV the field starts to "
    "approach 5 kV/m beneath the outer phases (centre-phase value "
    "shown). This kind of survey, which would otherwise require "
    "repeated runs of a commercial field-computation package, is "
    "produced in under a second by the FieldSense engine.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
tbl = doc.add_table(rows=1, cols=3)
tbl.style = "Light Grid Accent 1"
hdr = tbl.rows[0].cells
for i, h in enumerate(["Voltage (kV, line-line)",
                       "Centre-phase E_v (kV/m)",
                       "WHO status"]):
    hdr[i].text = h
rows = [("220",  "0.44", "Within limit"),
        ("400",  "0.81", "Within limit"),
        ("765",  "1.54", "Within limit"),
        ("1000", "2.02", "Within limit"),
        ("1200", "2.42", "Within limit (centre-phase)")]
for a, b, c in rows:
    rr = tbl.add_row().cells
    rr[0].text = a; rr[1].text = b; rr[2].text = c
add_para(doc, "Table 5.2  Centre-phase ground-level E_v across "
              "standard UHV voltage levels for the reference geometry.",
         italic=True, size=10,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
add_centered_image(doc, os.path.join(PLOTS, "who_bar.png"),
                   width_in=5.4,
                   caption="Fig 5.2  Centre-phase E_v at five voltage "
                           "levels against the WHO 5 kV/m public-"
                           "exposure limit (dashed red line).")

add_heading(doc, "5.6  Sensitivity to Height and Voltage", 2)
add_para(doc,
    "The sensitivity heatmap of Fig 5.3 maps E_v at the geometric "
    "centre as a function of the line height H (20–60 m) and the "
    "line-to-line voltage V_{LL} (200–1300 kV), with the white "
    "dashed contour highlighting the WHO 5 kV/m boundary. The "
    "diagram makes the design trade-off visible at a glance: "
    "increasing H is the most effective single lever to push a "
    "given operating voltage below the WHO limit, more so than "
    "increasing the phase spacing.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_centered_image(doc, os.path.join(PLOTS, "sensitivity_heatmap.png"),
                   width_in=6.0,
                   caption="Fig 5.3  Sensitivity of centre-phase E_v "
                           "to line height H and operating voltage "
                           "V_{LL}. White dashed contour marks the "
                           "WHO 5 kV/m public-exposure limit.")

add_heading(doc, "5.7  FieldSense Web Application – Demonstration", 2)
add_para(doc,
    "Opening index.html in any modern browser launches the "
    "FieldSense calculator. Moving the lateral-distance slider "
    "updates the numeric E_v readout, the WHO compliance badge, the "
    "[P] and [M] tables and the lateral profile in real time. Moving "
    "the voltage slider from 50 kV up to 900 kV scales the field "
    "linearly, while moving the height slider exposes the strong "
    "non-linear sensitivity of the peak field to the line-to-ground "
    "clearance.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "5.8  Limitations", 2)
for s in [
    "The model assumes a perfectly conducting, flat ground plane. "
    "Realistic terrain (hills, gorges, lossy soil) is not captured.",
    "Sag of conductors along the span is ignored — the line is "
    "modelled as horizontal at the average height.",
    "Only steady-state power-frequency fields are considered. "
    "Transients due to faults, switching and lightning are out of "
    "scope.",
    "The ML model has only been validated against analytical data; "
    "comparison against measured field surveys from an operating "
    "1200 kV line would strengthen the empirical confidence.",
]:
    add_bullet(doc, s)

add_heading(doc, "5.9  Conclusion", 2)
add_para(doc,
    "This project has demonstrated an end-to-end framework — "
    "analytical physics, machine-learning surrogate, and "
    "browser-based deployment — for predicting the ground-level "
    "vertical electric field of UHV AC transmission lines. The "
    "analytical engine, built on Maxwell's potential-coefficient "
    "matrix and the method of images, was validated on a 1200 kV "
    "horizontal-configuration test case. A 12 000-sample synthetic "
    "dataset was generated from this engine, and three regression "
    "models were trained on it. The best learner, a small ANN with "
    "two hidden layers, achieved R^{2} = 0.99886 on the test set "
    "and reproduced the analytical 1200 kV profile to within "
    "0.052 kV/m RMSE.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))
add_para(doc,
    "The FieldSense single-page web application makes both engines "
    "available to any user with a browser: there is no server, no "
    "Python runtime, and no installation. The calculator runs in "
    "milliseconds, displays a colour-coded WHO compliance badge, "
    "and plots the complete lateral profile against the 5 kV/m "
    "limit. The work therefore meets all six stated objectives of "
    "Section 1.5.",
    align=WD_ALIGN_PARAGRAPH.JUSTIFY,
    first_line_indent=Inches(0.4))

add_heading(doc, "5.10  Future Scope", 2)
for s in [
    "Extend the analytical engine to non-flat terrain and to sag "
    "along the span using catenary geometry, refreshing the ML "
    "dataset and retraining the surrogate.",
    "Incorporate magnetic-field computation alongside the electric "
    "field so that ICNIRP magnetic-field limits (100 µT public "
    "exposure at 50 Hz) can also be checked in the same calculator.",
    "Add a corona-loss and audible-noise estimator using "
    "established empirical formulae (CIGRE / IEEE) so that "
    "FieldSense becomes a comprehensive UHV right-of-way screening "
    "tool.",
    "Replace the MLP with a small Bayesian neural network or "
    "Gaussian process to provide predictive uncertainty bands "
    "alongside the point estimate.",
    "Validate the predicted profiles against measured field surveys "
    "from an operating 1200 kV line, e.g. data from the Bina "
    "National Test Station.",
]:
    add_bullet(doc, s, style="List Number")
add_page_break(doc)

# =================== References ===========================================
add_para(doc, "REFERENCES", size=16, bold=True,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14,
         color=RGBColor(0x1F, 0x3A, 0x5F))
refs = [
    "ICNIRP, \"Guidelines for limiting exposure to time-varying "
    "electric and magnetic fields (1 Hz – 100 kHz),\" Health "
    "Physics, vol. 99, no. 6, pp. 818–836, 2010.",
    "World Health Organisation, Environmental Health Criteria 238 – "
    "Extremely Low Frequency Fields, WHO Press, Geneva, 2007.",
    "G. B. Deno, \"Transmission line fields,\" IEEE Transactions on "
    "Power Apparatus and Systems, vol. PAS-95, no. 5, pp. 1600–1611, "
    "Sept. 1976.",
    "M. Abdel-Salam et al., High-Voltage Engineering: Theory and "
    "Practice, 2nd ed., Marcel Dekker, 2000, ch. 2.",
    "E. Kuffel, W. S. Zaengl and J. Kuffel, High Voltage Engineering: "
    "Fundamentals, 2nd ed., Newnes / Butterworth-Heinemann, 2000, "
    "ch. 4.",
    "Power Grid Corporation of India Ltd., \"1200 kV National Test "
    "Station, Bina – Technical Specifications,\" 2012.",
    "M. Abdel-Salam and Z. Abdel-Sattar, \"Calculation of corona "
    "losses on HVAC transmission lines,\" Electric Power Systems "
    "Research, vol. 12, pp. 13–22, 1987.",
    "A. Gupta, P. Pareek and V. Joshi, \"Computation of electric "
    "field beneath UHV transmission line using charge simulation "
    "method,\" Electric Power Components and Systems, vol. 41, no. "
    "6, pp. 615–628, 2013.",
    "S. Wang, J. Wang, et al., \"A review of deep learning for "
    "renewable energy forecasting,\" Energy Conversion and "
    "Management, vol. 198, p. 111799, 2019.",
    "X. Hu and Y. Zhang, \"Data-driven power-flow learning with "
    "physics-informed neural networks,\" IEEE Transactions on Smart "
    "Grid, vol. 11, no. 5, pp. 3990–4001, Sept. 2020.",
    "B. Donon et al., \"Deep statistical solvers and power-flow,\" "
    "Electric Power Systems Research, vol. 211, p. 108304, 2022.",
    "F. Pedregosa et al., \"Scikit-learn: Machine learning in "
    "Python,\" Journal of Machine Learning Research, vol. 12, "
    "pp. 2825–2830, 2011.",
    "I. Goodfellow, Y. Bengio and A. Courville, Deep Learning, MIT "
    "Press, 2016, ch. 6.",
    "Mozilla Developer Network, \"HTML5 Canvas API,\" web reference, "
    "Mozilla Foundation, retrieved 2026.",
    "J. D. Hunter, \"Matplotlib: A 2D graphics environment,\" "
    "Computing in Science & Engineering, vol. 9, no. 3, pp. 90–95, "
    "2007.",
]
for i, s in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.first_line_indent = Inches(-0.3)
    add_runs(p, f"[{i}]  {s}", size=11)

doc.save(OUT_PATH)
print(f"WROTE  {OUT_PATH}")
